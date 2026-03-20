"""
Generate ADBP_System_Rules_v3.docx from the markdown content.
Clean formatting matching the original ADBP_Final_Model_Legal style.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)


# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(text, level):
    p = doc.add_heading(text, level=level)
    p.style.font.name = "Calibri"
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def body(text, bold_spans=None):
    """Add a body paragraph. bold_spans = list of substrings to bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_spans:
        remaining = text
        for span in bold_spans:
            idx = remaining.find(span)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                set_font(r, 10)
            r = p.add_run(span)
            set_font(r, 10, bold=True)
            remaining = remaining[idx + len(span) :]
        if remaining:
            r = p.add_run(remaining)
            set_font(r, 10)
    else:
        r = p.add_run(text)
        set_font(r, 10)
    return p


def bullet(text, bold_spans=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_spans:
        remaining = text
        for span in bold_spans:
            idx = remaining.find(span)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                set_font(r, 10)
            r = p.add_run(span)
            set_font(r, 10, bold=True)
            remaining = remaining[idx + len(span) :]
        if remaining:
            r = p.add_run(remaining)
            set_font(r, 10)
    else:
        r = p.add_run(text)
        set_font(r, 10)
    return p


def numbered(text, num=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    set_font(r, 10)
    return p


def blockquote(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, 10, bold=True, italic=True)
        r = p.add_run(text[len(bold_prefix) :])
        set_font(r, 10, italic=True)
    else:
        r = p.add_run(text)
        set_font(r, 10, italic=True)
    # Left border shading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "AAAAAA")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def code_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 150)
    return p


def add_table(headers, rows, col_widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        # Gray header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Inches(col_widths[ci])
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# Title / header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("American Dream Benefits Program | System Rules v3")
set_font(r, 16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Confidential | March 2026")
set_font(r, 9, color=(100, 100, 100))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Status: Updated to reflect confirmed model mechanics")
set_font(r, 9, color=(100, 100, 100))

hr()

# ── Introduction ──────────────────────────────────────────────────────────────
heading("Introduction", 2)

body(
    "Goal Restore the American Dream — enable families to thrive on one job / one income again, "
    "without needing two incomes, debt, or constant financial pressure.",
    bold_spans=["Goal"],
)

body(
    "The American Dream Benefits Program is a third-party consumer discount platform with voluntary "
    "payroll facilitation that delivers substantial purchasing power for essentials."
)

body(
    "Core Mechanism A closed-loop discount-credit system powered by Solana blockchain. "
    "Employees voluntarily purchase Discount Credits (utility tokens) with after-tax dollars.",
    bold_spans=["Core Mechanism"],
)

blockquote(
    "$1 = 1 credit = $2 in purchasing power at participating essential vendors (50% discount)"
)

hr()

# ── Program Participants ──────────────────────────────────────────────────────
heading("Program Participants", 2)

bullet(
    "Vendors — essential businesses (groceries, utilities, rent, gas, etc.) that participate, "
    "employ participating employees, accept & recirculate credits",
    bold_spans=["Vendors"],
)
bullet(
    "Employers — businesses that participate, employ participating employees, accept & recirculate credits",
    bold_spans=["Employers"],
)
bullet(
    "Employees — any person working for a participating employer who voluntarily participates",
    bold_spans=["Employees"],
)
bullet(
    "Program — operates the Solana smart contracts, treasury, and platform",
    bold_spans=["Program"],
)

hr()

# ── System Rules ──────────────────────────────────────────────────────────────
heading("System Rules", 2)

heading("Credit Purchase", 3)
bullet(
    "Employees purchase credits post-tax via embedded fiat on-ramp (bank card, ACH, Apple Pay)"
)
bullet(
    "$1.00 paid by employee = 1 credit minted = $2.00 purchasing power at vendors",
    bold_spans=[
        "$1.00 paid by employee = 1 credit minted = $2.00 purchasing power at vendors"
    ],
)
bullet(
    "The full $1.00 purchase price flows to the treasury — this is the primary funding mechanism",
    bold_spans=["$1.00 purchase price flows to the treasury"],
)
bullet(
    "Employees are limited to 5,000 credits per month",
    bold_spans=["5,000 credits per month"],
)
bullet(
    "No general cash redemption — credits are non-redeemable for cash and stay in the closed loop"
)

heading("Admin Revenue Pool", 3)
bullet(
    "A 10% fee on all credit purchases is collected as admin revenue",
    bold_spans=["10% fee"],
)
bullet("At $1.00/credit, this equals $0.10 per credit minted")
bullet(
    "This fee does not flow to the treasury — it is tracked and distributed separately",
    bold_spans=["not"],
)
bullet(
    "Distribution is pro-rata by recirculation share: the more credits a vendor or employer "
    "recirculates within the network, the larger their share of the admin revenue pool",
    bold_spans=["pro-rata by recirculation share"],
)

blockquote(
    "Why pro-rata? It directly incentivizes recirculation — participants who actively move credits "
    "through the network earn proportionally more. A vendor recirculating 10% of total credits "
    "earns 10% of the admin pool.",
    bold_prefix="Why pro-rata?",
)

heading("Recirculation Rules", 3)
bullet(
    "Vendors decide the maximum credits they accept in a 12-month period (their recirculation %)"
)
bullet(
    "Employers and vendors recirculate credits for any of their expenses except payroll and taxes "
    "— creating an accounting wash ($X foregone income offset by $X saved)",
    bold_spans=["except"],
)
bullet(
    "The recirculation-to-mint ratio is 2:1 — a vendor/employer that can recirculate 100,000 "
    "credits allows their employees to mint 50,000 total",
    bold_spans=["recirculation-to-mint ratio is 2:1"],
)
bullet(
    "Credits are automatically recirculated at month-end on behalf of vendors and employers"
)
bullet("Velocity is pegged at 12\u00d7 per year", bold_spans=["12\u00d7 per year"])

heading("Burn Mechanic (Discretionary)", 3)
bullet(
    "Burns are manual and discretionary — no automatic per-transaction burn",
    bold_spans=["manual and discretionary"],
)
bullet(
    "The treasury pays $2.00 per credit burned (the $2 fair market value of the credit)",
    bold_spans=["$2.00 per credit burned"],
)
bullet(
    "Burns are threshold-triggered: a burn is authorized when the backing ratio "
    "(treasury_wallet / total_credits) exceeds a target threshold"
)
bullet(
    "Affordability cap: post-burn backing ratio must never fall below the 50% floor",
    bold_spans=["Affordability cap:"],
)
code_line(
    "Cap formula: max_burnable = (wallet \u2212 0.50 \u00d7 total_credits) / (2.00 \u2212 0.50)"
)

bullet(
    "MC-optimal strategy (confirmed via 20,000-run Monte Carlo):",
    bold_spans=["MC-optimal strategy"],
)
bullet("Trigger: backing \u2265 114.6%", level=1)
bullet("Size: 30.2% of outstanding credits per event", level=1)
bullet("Cooldown: 8 months minimum between burns", level=1)
bullet("First eligible: month 13 (ramp-up protection)", level=1)
bullet(
    "Result: 1 burn event over 10 years, 30.2% of credits destroyed, 77.7% final backing (HEALTHY)",
    level=1,
)

blockquote(
    "Why $2/credit? The credit provides $2 of purchasing power. Burning a credit extinguishes "
    "a $2 obligation. The treasury pays $2 to retire $1 of face-value liability — the difference "
    "is the amplification cost built into the burn mechanism.",
    bold_prefix="Why $2/credit?",
)

heading("Treasury Mechanics", 3)
add_table(
    ["Flow", "Amount", "Direction"],
    [
        ["Credit purchase", "$1.00/credit", "\u2192 Treasury"],
        ["Interest income", "4% APR on wallet balance (monthly)", "\u2192 Treasury"],
        ["Burn event", "$2.00/credit burned", "\u2190 Treasury"],
        ["Admin fee", "$0.10/credit (10%)", "\u2192 Admin pool (separate)"],
    ],
    [2.3, 2.3, 1.9],
)
code_line("Interest timing: wallet[t] = wallet[t-1] + interest[t-1] + inflow[t]")
code_line("Then: interest[t] = wallet[t] \u00d7 (0.04 / 12)")

heading("Solvency Thresholds", 3)
add_table(
    ["Status", "Backing Ratio", "Definition"],
    [
        [
            "HEALTHY",
            "\u2265 75%",
            "Treasury holds \u2265 $0.75 per $1 of credit outstanding",
        ],
        ["WARNING", "50% \u2013 74%", "Solvent but under pressure"],
        ["FAILURE", "< 50%", "Treasury cannot cover basic obligations"],
    ],
    [1.1, 1.4, 4.0],
)
code_line("Primary metric: backing_ratio = treasury_wallet / total_credits_outstanding")
code_line(
    "Secondary metric: burn_coverage = treasury_wallet / (total_credits \u00d7 $2.00)"
)

hr()

# ── Recirculation Flywheel ────────────────────────────────────────────────────
heading("Recirculation Flywheel", 2)
for line in [
    "Employee buys credits \u2192 smart contract mints tokens \u2192 $1/credit to treasury \u2192 $0.10/credit to escrow",
    "Employee spends at vendors \u2192 50% discount applied (2\u00d7 amplification)",
    "Vendor receives tokens \u2192 recirculates to suppliers, utilities, etc.",
    "Loop continues at 12\u00d7 annual velocity",
    "No credits exit for cash, ever",
    "Admin revenue pool distributed monthly pro-rata by recirculation share from escrow wallet",
]:
    numbered(line)

# ── Essential Businesses ──────────────────────────────────────────────────────
heading("Essential Businesses (Vendors)", 2)
for vendor in [
    "Health Insurance Providers",
    "Gas Stations",
    "Grocery Stores",
    "Rental Housing",
    "Internet Service Provider",
    "Electric and Gas Utilities",
    "Water Utilities",
    "Childcare Facilities",
    "Car Service (Repairs)",
]:
    numbered(vendor)

hr()

# ── Compliance ────────────────────────────────────────────────────────────────
heading("Compliance & Regulatory Posture", 2)
for line in [
    "Non-custodial utility token on Solana (closed-loop discount access).",
    "No private currency risk under Stamp Payments Act (\u00a7 486) - decentralized, no centralized issuer. "
    "Credits are utility tokens in a closed-loop system, non-redeemable for cash, decentralized on Solana, "
    "and not intended as general currency. 2026 guidance on tokenized assets distinguishes utility tokens "
    "from scrip if they are platform-specific and non-convertible. The GENIUS Act (2025) for stable coins "
    "explicitly exempts non-currency digital assets, reinforcing this.",
    "No ERISA risk - third-party platform, not employer-sponsored. No welfare benefit triggers like health "
    "insurance or childcare. Vendors and employers earn a fair market value fixed admin fee for offering a "
    "discount and recirculating credits.",
    "No money transmission / CVC risk - MSB partner handles fiat entry. Per FinCEN\u2019s longstanding "
    "guidance, convertible virtual currency (CVC) is virtual currency that either has an equivalent value "
    "in real currency or acts as a substitute for it. Closed-loop systems where credits cannot exit for "
    "cash are explicitly distinguished from CVC.",
    "No private scrip risk - the Stamp Payments Act targets private currencies that compete with fiat, not "
    "closed-loop discounts. Decentralized, blockchain-based utility tokens without cash equivalence or broad "
    "circulation are generally exempt. Credits recirculate indefinitely but only for essentials among "
    "participants - no public circulation or fiat substitution outside the loop. The decentralized design "
    "(no central issuer) further reduces risk.",
    "Barter tax compliance - automated 1099-B reporting at $2 FMV per credit.",
    "Regional launch - NJ to start with regional expansion as the business grows.",
]:
    bullet(line)

# ── Bottom Line ───────────────────────────────────────────────────────────────
heading("Bottom Line", 2)
body(
    "One job. One family. One future. A win-win system that can turn a $60,000 salary into a potential "
    "$120,000 of essential purchasing power while financially incentivizing the vendor or employer - "
    "powered by Solana blockchain, modern UX, and built-in compliance."
)
body("Ready to launch regionally and scale responsibly.")

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "American Dream Benefits Program  |  Confidential \u2014 Attorney Review Copy  |  March 2026"
)
set_font(r, 8, color=(130, 130, 130))

doc.save("docs/ADBP_System_Rules_v3.docx")
print("Word doc written: docs/ADBP_System_Rules_v3.docx")
